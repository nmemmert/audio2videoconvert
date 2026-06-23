"""
Shared render pipeline for the podcast video generator.
Call render_job(settings, audio_path, output_path, ...) to produce a video.
"""

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
from pathlib import Path
from string import Template

try:
    from docx import Document
except ImportError:
    Document = None

HERE = Path(__file__).resolve().parent

DEFAULTS = {
    # Art
    "art_enabled": True,
    "art": "podcast-art.jpeg",
    "art_size": 460,
    "art_x_offset": 0,
    # Glow
    "glow_enabled": True,
    "glow_color": "#c9a84c",
    "glow_sigma": 40,
    # Background
    "bg_style": "solid",
    "bg_color": "#111111",
    "bg_color2": "#1a1a2e",
    # Waveform animation
    "waveform_enabled": True,
    "waveform_color": "#c9a84c",       # mid-amplitude (normal waves)
    "waveform_high_color": "#cc3333",  # peaks / loud moments
    "waveform_low_color": "#3366cc",   # quiet / low amplitude
    "waveform_height": 100,
    # Captions
    "font_name": "Georgia",
    "font_size": 64,
    "words_per_chunk": 2,
    "caption_color": "&H50D8EAF0",
    "caption_highlight": "&H004CA8C9",
    "caption_back": "&HBF000000",
    "caption_style": "karaoke",
    # Outline
    "outline_enabled": False,
    "outline_style": "sidebar",  # "sidebar" or "ticker"
    "outline_color": "#ffffff",
    "outline_font": "Georgia",
    "outline_font_size": 32,
    # Title overlay
    "title_enabled": False,
    "title_font": "Georgia",
    "title_font_size": 48,
    "title_color": "white",
    # Output
    "width": 1920,
    "height": 1080,
    "fps": 30,
    "output_dir": "",
    # Discussion question cards
    "question_cards_enabled": False,
    # QR code / watermark
    "qr_enabled": False,
    "qr_path": "",
    "qr_size": 160,
    "qr_corner": "bottom-right",  # bottom-right, bottom-left, top-right, top-left
    # Ollama (outline generation fallback)
    "ollama_url": "http://localhost:11434",
    "ollama_model": "llama3.1",
}

BG_STYLES = ["solid", "gradient"]
CAPTION_STYLES = ["karaoke", "fade", "bounce"]
OUTLINE_STYLES = ["sidebar", "ticker"]

PYTHON_VENV = Path.home() / "whisper-env" / "bin" / "python3"


def find_python():
    if PYTHON_VENV.exists():
        return str(PYTHON_VENV)
    return sys.executable


def find_ffmpeg(name="ffmpeg"):
    p = shutil.which(name)
    return p or name


def detect_glow_color(art_path, fallback="#c9a84c"):
    try:
        from PIL import Image
        img = Image.open(art_path).convert("RGB")
        w, h = img.size
        m = max(2, min(w, h) // 50)
        corners = [(m, m), (w - m, m), (m, h - m), (w - m, h - m)]
        r, g, b = 0, 0, 0
        for x, y in corners:
            pr, pg, pb = img.getpixel((x, y))
            r += pr; g += pg; b += pb
        n = len(corners)
        r, g, b = r // n, g // n, b // n
        if max(r, g, b) < 24:
            return fallback
        return f"#{r:02x}{g:02x}{b:02x}"
    except Exception:
        return fallback


def hex_to_ass(hex_color, alpha="00"):
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        hex_color = "ffffff"
    rr, gg, bb = hex_color[0:2], hex_color[2:4], hex_color[4:6]
    return f"&H{alpha}{bb}{gg}{rr}".upper()


# ── Outline / transcript helpers ────────────────────────────────────────────

_HEADER_RE = re.compile(r"^(?:[^A-Za-z0-9]+\s*)?([A-Z][A-Za-z0-9 ,:;'\-—]+)$")
_SPEAKER_RE = re.compile(r"^[A-Z][A-Z ]*:\s*")


def _normalize(text):
    text = text.lower()
    text = re.sub(r"[^a-z0-9 ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def extract_discussion_questions(docx_path):
    """Find discussion questions in a docx script.
    Returns list of (question_text, anchor_text) suitable for timestamp alignment."""
    if Document is None:
        return []
    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    questions = []
    in_discussion = False

    for text in paragraphs:
        if not text:
            continue
        upper = text.upper()
        # Enter a discussion/question block
        if "DISCUSSION" in upper or "SMALL GROUP" in upper or "STUDY QUESTION" in upper:
            in_discussion = True
            continue
        # Exit on a new all-caps section header that isn't a question
        if in_discussion and _HEADER_RE.match(text) and "DISCUSSION" not in upper and "QUESTION" not in upper:
            if text == text.upper() or text.startswith("SEGMENT") or text.startswith("CLOSING"):
                in_discussion = False
                continue
        if in_discussion:
            # Strip leading "1." / "2)" / "Q1." numbering
            clean = re.sub(r"^[Q\d]+[\.\)]\s*", "", text).strip()
            if clean and (clean.endswith("?") or re.match(r"^[Q\d]+[\.\)]", text)):
                questions.append((clean, clean))

    return questions


def extract_outline_from_script(docx_path):
    if Document is None:
        return []
    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    points = []
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i]
        m = _HEADER_RE.match(text) if text else None
        if m:
            title = m.group(1).strip()
            anchor = ""
            for j in range(i + 1, len(paragraphs)):
                cand = paragraphs[j].strip()
                if not cand:
                    continue
                if _HEADER_RE.match(cand):
                    break
                if cand.startswith("[") and cand.endswith("]"):
                    continue
                if set(cand) <= set("— ✦-—  "):
                    continue
                anchor = _SPEAKER_RE.sub("", cand)
                break
            points.append((title, anchor))
        i += 1
    return points


def align_outline_to_transcript(points, segments):
    if not points or not segments:
        return []
    seg_norms = [_normalize(seg["text"]) for seg in segments]
    offsets, full_parts, pos = [], [], 0
    for norm in seg_norms:
        offsets.append(pos)
        full_parts.append(norm)
        pos += len(norm) + 1
    full_text = " ".join(full_parts)
    result, search_from, last_time = [], 0, 0.0
    for title, anchor in points:
        anchor_norm = _normalize(anchor)
        words = anchor_norm.split()
        idx = -1
        for n in (6, 4, 3):
            if len(words) >= n:
                needle = " ".join(words[:n])
                idx = full_text.find(needle, search_from)
                if idx != -1:
                    break
        if idx == -1:
            continue
        seg_idx = 0
        for k, off in enumerate(offsets):
            if off <= idx:
                seg_idx = k
            else:
                break
        t = max(segments[seg_idx]["start"], last_time)
        result.append((t, title))
        last_time = t
        search_from = idx + 1
    return result


def generate_outline(segments, s, log_cb=None):
    import urllib.request

    def log(msg):
        if log_cb:
            log_cb(msg)

    transcript_lines = "\n".join(f"[{seg['start']:.1f}s] {seg['text']}" for seg in segments)
    prompt = (
        "You are given a transcript of a podcast episode, where each line is "
        "prefixed with the timestamp (in seconds) at which it was spoken.\n\n"
        "Produce a short outline of the main points/topics discussed, in order. "
        "For each point, pick the timestamp where that topic begins.\n\n"
        "Respond ONLY with a JSON array of objects, each with keys \"time\" (number, "
        "seconds) and \"text\" (string, max ~6 words). Aim for 4-8 points total.\n\n"
        f"Transcript:\n{transcript_lines}\n"
    )
    url = s.get("ollama_url", "http://localhost:11434").rstrip("/") + "/api/generate"
    payload = json.dumps({
        "model": s.get("ollama_model", "llama3.1"),
        "prompt": prompt,
        "stream": False,
        "format": "json",
    }).encode("utf-8")
    try:
        req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        points = json.loads(data.get("response", ""))
        result = [(float(p["time"]), str(p["text"]).strip()) for p in points if p.get("text")]
        result.sort(key=lambda x: x[0])
        return result
    except Exception as e:
        log(f"Outline generation failed: {e}")
        return []


def _ass_fmt(t):
    h = int(t // 3600)
    m = int((t % 3600) // 60)
    sec = t % 60
    return f"{h:d}:{m:02d}:{sec:05.2f}"


def _subtract_intervals(start, end, blackouts):
    """Return sub-intervals of [start, end] that don't overlap any blackout range."""
    result = []
    cur = start
    for bs, be in sorted(blackouts):
        if be <= cur:
            continue
        if bs >= end:
            break
        if bs > cur:
            result.append((cur, bs))
        cur = max(cur, be)
    if cur < end:
        result.append((cur, end))
    return result


def build_sidebar_ass(points, s, duration_f, out_path, start_y=None, end_y=None,
                      art_x=None, blackout=None):
    """Left-side outline: one item per line, expanding to fill the sidebar width.

    Font size is the largest value where:
      (a) even the longest item fits on a single line within text_px_w, AND
      (b) all items fit vertically within start_y → end_y

    No Python word-wrapping — each item is always a single line so the text
    fills right up to the art edge for every item.

    start_y  — top of the outline region (below captions)
    end_y    — bottom of the outline region (above waveform)
    art_x    — left edge of the art image; text stays left of this
    blackout — list of (start, end) intervals where the outline is suppressed
    """
    blackout = blackout or []

    color = hex_to_ass(s.get("outline_color", "#ffffff"))
    font = s.get("outline_font", "Georgia")
    wave_h = s.get("waveform_height", 100) if s.get("waveform_enabled") else 0
    height = s["height"]

    if start_y is None:
        start_y = 24
    if end_y is None:
        end_y = height - wave_h - 12
    if art_x is None:
        art_x = s["width"] // 2

    n = max(len(points), 1)
    available_h = max(1, end_y - start_y)
    # Usable pixel width: x=24 to art_x-24
    text_px_w = max(80, art_x - 48)

    # Longest item drives the horizontal font cap
    longest = max((len(text) for _, text in points), default=1)
    # Georgia character width ≈ size × 0.62 px; be slightly conservative at 0.65
    size_for_width  = max(18, int(text_px_w / (longest * 0.65)))
    # Each item gets equal vertical slot; font must fit within slot
    size_for_height = max(18, (available_h // n) - 10)
    size = min(52, size_for_width, size_for_height)
    size = max(18, size)

    # Equal slot per item; 72% used for text, 28% is breathing room
    item_slot = available_h / n
    FILL = 0.72

    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: {s['width']}
PlayResY: {s['height']}
WrapStyle: 0

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,{font},{size},{color},{color},&H00000000,&H80000000,0,0,0,0,100,100,1,0,1,2,0,7,24,24,24,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    lines = [header]
    dim_hex = "888888"
    line_h = size  # single line per item — no wrapping
    for i, (start, text) in enumerate(points):
        escaped = text.replace("\\", "\\\\").replace("{", "").replace("}", "")
        slot_top = start_y + i * item_slot
        margin   = item_slot * (1 - FILL) / 2
        y = int(slot_top + margin + (item_slot * FILL - line_h) / 2)

        if start < duration_f:
            pos = f"\\pos(24,{y})\\an7"
            next_start = points[i + 1][0] if i + 1 < len(points) else duration_f
            active_end = min(next_start, duration_f)
            for seg_s, seg_e in _subtract_intervals(start, active_end, blackout):
                tag = f"{{{pos}\\fad(400,0)}}"
                lines.append(f"Dialogue: 0,{_ass_fmt(seg_s)},{_ass_fmt(seg_e)},Default,,0,0,0,,{tag}{escaped}")
            for seg_s, seg_e in _subtract_intervals(active_end, duration_f, blackout):
                tag = f"{{{pos}\\c&H{dim_hex}&}}"
                lines.append(f"Dialogue: 0,{_ass_fmt(seg_s)},{_ass_fmt(seg_e)},Default,,0,0,0,,{tag}{escaped}")

    out_path.write_text("\n".join(lines))


def build_ticker_ass(points, s, duration_f, out_path):
    """Bottom chapter ticker: show current chapter name in lower-left."""
    color = hex_to_ass(s.get("outline_color", "#ffffff"))
    font = s.get("outline_font", "Georgia")
    size = s.get("outline_font_size", 32)
    wave_h = s.get("waveform_height", 100) if s.get("waveform_enabled") else 0
    y = s["height"] - wave_h - size - 16

    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: {s['width']}
PlayResY: {s['height']}
WrapStyle: 0

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,{font},{size},{color},{color},&H00000000,&HA0000000,0,0,0,0,100,100,1,0,4,2,0,1,30,30,{y},1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    lines = [header]
    for i, (start, text) in enumerate(points):
        if start >= duration_f:
            continue
        end = points[i + 1][0] if i + 1 < len(points) else duration_f
        escaped = text.replace("\\", "\\\\").replace("{", "").replace("}", "")
        tag = "{\\fad(300,300)}"
        lines.append(f"Dialogue: 0,{_ass_fmt(start)},{_ass_fmt(end)},Default,,0,0,0,,{tag}{escaped}")
    out_path.write_text("\n".join(lines))


def align_questions_to_transcript(questions, segments):
    """Find when each discussion question is introduced in the audio.

    Strategy (tried in order for each question):
    1. Search transcript for 'question' + ordinal/number near that position
       (e.g. host says "question one", "question 1", "first question")
    2. Fallback: search for first 3 words of the question text verbatim

    Each question is searched only after the previous one was found, so
    ordering is always preserved.
    """
    if not questions or not segments:
        return []

    _ORDINALS = [
        ["first", "one", "1"],
        ["second", "two", "2"],
        ["third", "three", "3"],
        ["fourth", "four", "4"],
        ["fifth", "five", "5"],
        ["sixth", "six", "6"],
        ["seventh", "seven", "7"],
        ["eighth", "eight", "8"],
    ]

    result = []
    search_from_seg = 0  # don't re-scan segments before last found question

    for i, (q_text, _) in enumerate(questions):
        markers = _ORDINALS[i] if i < len(_ORDINALS) else [str(i + 1)]
        found_t = None

        # Pass 1: look for "question number [n]", "question [ordinal]", etc.
        number_str = str(i + 1)
        all_markers = markers + [number_str]
        for si in range(search_from_seg, len(segments)):
            seg_lower = segments[si]["text"].lower()
            if "question" in seg_lower:
                for m in all_markers:
                    # matches "question number 1", "question 1", "question one", etc.
                    if m in seg_lower:
                        found_t = segments[si]["start"]
                        search_from_seg = si + 1
                        break
            if found_t is not None:
                break

        # Pass 2: match first 3 words of question text verbatim
        if found_t is None:
            q_words = _normalize(q_text).split()
            needle = " ".join(q_words[:3]) if len(q_words) >= 3 else _normalize(q_text)
            for si in range(search_from_seg, len(segments)):
                if needle in _normalize(segments[si]["text"]):
                    found_t = segments[si]["start"]
                    search_from_seg = si + 1
                    break

        if found_t is not None:
            result.append((found_t, q_text))

    return result


def build_question_drawbox(timed_questions, duration_f, card_x, card_y, card_w, card_h):
    """Return an ffmpeg drawbox filter string that covers the right panel during each question.
    The box snaps in/out (instant); the ASS text layer handles the fade."""
    if not timed_questions:
        return ""
    parts = []
    for i, (start, text) in enumerate(timed_questions):
        if start >= duration_f:
            continue
        next_start = timed_questions[i + 1][0] if i + 1 < len(timed_questions) else duration_f
        end = min(next_start, start + 55, duration_f)
        if end <= start:
            continue
        parts.append(f"between(t,{start:.3f},{end:.3f})")
    if not parts:
        return ""
    enable = "gt(" + "+".join(parts) + ",0)"
    return (f"drawbox=x={card_x}:y={card_y}:w={card_w}:h={card_h}"
            f":color=0x0a0a0a@1.0:t=fill:enable='{enable}'")


def build_question_cards_ass(timed_questions, s, duration_f, out_path,
                              card_x, card_y, card_w, card_h,
                              text_cx=None, text_cy=None):
    """ASS text overlay for discussion question cards.
    The dark background is handled separately by build_question_drawbox().
    Text is centred at (text_cx, text_cy) and fades in/out.
    """
    import textwrap as _tw

    font = s.get("outline_font", "Georgia")
    width, height = s["width"], s["height"]
    card_cx = text_cx if text_cx is not None else card_x + card_w // 2
    card_cy = text_cy if text_cy is not None else card_y + card_h // 2

    pad = 60
    # Text area width = the usable space on each side of card_cx, doubled.
    # This prevents text from overflowing left into the outline or right off screen.
    half_w = min(card_cx - card_x, (card_x + card_w) - card_cx) - pad
    text_w_px = max(200, half_w * 2)
    # Font size based on the actual text area width, not the full card
    font_size = max(26, min(48, text_w_px // 20))
    chars_per_line = max(12, int(text_w_px / (font_size * 0.54)))

    header = f"""[Script Info]
ScriptType: v4.00+
PlayResX: {width}
PlayResY: {height}
WrapStyle: 2

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Text,{font},{font_size},&H00FFFFFF,&H00FFFFFF,&H00000000,&H00000000,0,0,0,0,100,100,2,0,1,1,0,5,0,0,0,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""
    lines = [header]

    for i, (start, text) in enumerate(timed_questions):
        if start >= duration_f:
            continue
        next_start = timed_questions[i + 1][0] if i + 1 < len(timed_questions) else duration_f
        end = min(next_start, start + 55, duration_f)
        if end <= start:
            continue

        s_fmt, e_fmt = _ass_fmt(start), _ass_fmt(end)
        escaped = text.replace("\\", "\\\\").replace("{", "").replace("}", "")
        wrapped = _tw.wrap(escaped, chars_per_line) or [escaped]
        ass_text = r"\N".join(wrapped)
        tag = f"{{\\an5\\pos({card_cx},{card_cy})\\fad(700,700)}}"
        lines.append(f"Dialogue: 0,{s_fmt},{e_fmt},Text,,0,0,0,,{tag}{ass_text}")

    out_path.write_text("\n".join(lines))


def build_bg_filter(s, fps):
    width, height, bg = s["width"], s["height"], s["bg_color"]
    if s.get("bg_style") == "gradient":
        c2 = s.get("bg_color2", "#1a1a2e")
        return (
            f"gradients=s={width}x{height}:r={fps}:c0={bg}:c1={c2}:"
            f"x0=0:y0=0:x1=0:y1={height}[bg]"
        )
    return f"color=c={bg}:s={width}x{height}:r={fps}[bg]"


TRANSCRIBE_TEMPLATE = Template('''
import warnings
warnings.filterwarnings("ignore")
from faster_whisper import WhisperModel

model = WhisperModel("base", device="cpu", compute_type="int8")
segments, info = model.transcribe(r"$AUDIO", word_timestamps=True)

CHUNK_SIZE = $CHUNK_SIZE
CAPTION_STYLE = "$CAPTION_STYLE"

def fmt_time(t):
    h = int(t // 3600)
    m = int((t % 3600) // 60)
    s = t % 60
    return f"{h:d}:{m:02d}:{s:05.2f}"

ass_header = """[Script Info]
ScriptType: v4.00+
PlayResX: $WIDTH
PlayResY: $HEIGHT
WrapStyle: 2

[V4+ Styles]
Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding
Style: Default,$FONT_NAME,$FONT_SIZE,$CAPTION_HIGHLIGHT,$CAPTION_COLOR,&H00000000,$CAPTION_BACK,-1,0,0,0,100,100,1,0,4,0,0,2,60,60,$CAPTION_Y,1

[Events]
Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text
"""

lines = [ass_header]
import json as _json
all_words = []
seg_list = []
for segment in segments:
    seg_list.append({"start": segment.start, "end": segment.end, "text": segment.text.strip()})
    words = list(segment.words) if segment.words else []
    if not words:
        class FakeWord:
            pass
        w = FakeWord()
        w.word = segment.text.strip()
        w.start = segment.start
        w.end = segment.end
        all_words.append(w)
    else:
        for w in words:
            w.word = w.word.strip()
            all_words.append(w)
    print(f"PROGRESS {segment.end:.2f}", flush=True)

with open(r"$SEGMENTS_PATH", "w") as f:
    _json.dump(seg_list, f)

if CAPTION_STYLE == "bounce":
    for w in all_words:
        if not w.word:
            continue
        text = "{\\\\fscx60\\\\fscy60\\\\t(0,120,\\\\fscx100\\\\fscy100)}" + w.word
        lines.append(f"Dialogue: 0,{fmt_time(w.start)},{fmt_time(w.end)},Default,,0,0,0,,{text}")
else:
    chunks = [all_words[i:i+CHUNK_SIZE] for i in range(0, len(all_words), CHUNK_SIZE)]
    for chunk in chunks:
        if not chunk:
            continue
        chunk_start = fmt_time(chunk[0].start)
        chunk_end = fmt_time(chunk[-1].end)
        if CAPTION_STYLE == "fade":
            text = "{\\\\fad(150,150)}" + " ".join(w.word for w in chunk)
        else:
            parts = []
            for w in chunk:
                duration_cs = max(1, int((w.end - w.start) * 100))
                parts.append("{\\\\k" + str(duration_cs) + "}" + w.word.strip())
            text = " ".join(parts)
        lines.append(f"Dialogue: 0,{chunk_start},{chunk_end},Default,,0,0,0,,{text}")

with open(r"$ASS_PATH", "w") as f:
    f.write("\\n".join(lines))

print("DONE")
''')


def _run_ffmpeg_with_progress(cmd, total_seconds, progress_cb):
    cmd = cmd + ["-progress", "pipe:1", "-nostats"]
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                             text=True, bufsize=1)
    stderr_chunks = []

    def drain_stderr():
        for chunk in proc.stderr:
            stderr_chunks.append(chunk)

    t = threading.Thread(target=drain_stderr, daemon=True)
    t.start()
    for line in proc.stdout:
        line = line.strip()
        if line.startswith("out_time_ms="):
            try:
                seconds = int(line.split("=", 1)[1]) / 1_000_000
                frac = min(1.0, seconds / total_seconds) if total_seconds else 0
                if progress_cb:
                    progress_cb(frac)
            except (ValueError, ZeroDivisionError):
                pass
    proc.wait()
    t.join()
    return proc.returncode, "".join(stderr_chunks)


def _run_transcription_with_progress(python_bin, env, script, total_seconds, progress_cb):
    proc = subprocess.Popen([python_bin, "-W", "ignore", "-c", script],
                             stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                             text=True, bufsize=1, env=env)
    stderr_chunks = []

    def drain_stderr():
        for chunk in proc.stderr:
            stderr_chunks.append(chunk)

    t = threading.Thread(target=drain_stderr, daemon=True)
    t.start()
    for line in proc.stdout:
        line = line.strip()
        if line.startswith("PROGRESS "):
            try:
                seconds = float(line.split(" ", 1)[1])
                frac = min(1.0, seconds / total_seconds) if total_seconds else 0
                if progress_cb:
                    progress_cb(frac)
            except (ValueError, ZeroDivisionError):
                pass
    proc.wait()
    t.join()
    return proc.returncode, "".join(stderr_chunks)


def _probe_duration(ffprobe, src):
    r = subprocess.run(
        [ffprobe, "-v", "error", "-show_entries", "format=duration",
         "-of", "csv=p=0", str(src)],
        capture_output=True, text=True,
    )
    try:
        return float(r.stdout.strip())
    except ValueError:
        return 0


def render_job(s, audio_path, output_path, art_path=None,
               title="", script_path=None,
               progress_cb=None, log_cb=None,
               preview_seconds=None, outline_titles=None):
    """
    Render one podcast video.

    s               : settings dict (see DEFAULTS)
    audio_path      : Path to source audio
    output_path     : Path for the final .mp4
    art_path        : Path to cover art image (None if art disabled)
    title           : optional episode title overlay text
    script_path     : optional .docx episode script for outline
    progress_cb     : callable(stage: str, fraction: float)
    log_cb          : callable(str)
    preview_seconds : render only this many seconds (quick sanity check)
    outline_titles  : optional list of replacement title strings
    """
    def log(msg):
        if log_cb:
            log_cb(msg)

    def progress(stage, frac):
        if progress_cb:
            progress_cb(stage, frac)

    ffmpeg = find_ffmpeg()
    ffprobe = find_ffmpeg("ffprobe")
    python_bin = find_python()

    audio = Path(audio_path)
    output = Path(output_path)

    art_enabled = s.get("art_enabled", True) and art_path is not None
    if art_enabled:
        art_path = Path(art_path)
        if not art_path.exists():
            art_enabled = False

    width = s["width"]
    height = s["height"]
    fps = s["fps"]
    art_size = s.get("art_size", 460)
    wave_enabled = s.get("waveform_enabled", True)
    wave_h = s.get("waveform_height", 100) if wave_enabled else 0
    outline_enabled = s.get("outline_enabled", False)
    outline_style = s.get("outline_style", "sidebar")

    # sidebar_w is only finalised after outline points are found.
    # We track the *requested* value here; art_x is recomputed after outline building.
    _sidebar_w_req = 480 if (outline_enabled and outline_style == "sidebar") else 0

    # Vertical layout — does not depend on sidebar_w, so compute it now.
    usable_h = height - wave_h
    caption_reserve = 90  # pixels above art reserved for captions
    art_y = max(caption_reserve, (usable_h - art_size) // 2)
    art_y = max(caption_reserve, min(art_y, usable_h - art_size - 10))
    art_y = max(0, art_y)

    # Caption sits above the art — computed from art_y which is now fixed.
    caption_y_ass = height - art_y + 16

    def _compute_art_x(sw):
        """Return clamped art_x for a given sidebar width sw."""
        ax = sw + (width - sw - art_size) // 2 + s.get("art_x_offset", 0)
        return max(0, min(ax, width - art_size))

    # Preliminary art_x (used for outline text-wrap if points are found).
    # Will be recomputed after outline building with the true sidebar_w.
    sidebar_w = _sidebar_w_req
    art_x = _compute_art_x(sidebar_w)

    glow_color = detect_glow_color(art_path, s.get("glow_color", "#c9a84c")) if art_enabled else s.get("glow_color", "#c9a84c")

    ass_path = None
    outline_ass_path = None
    preview_audio = None

    try:
        if preview_seconds:
            preview_audio = Path(tempfile.gettempdir()) / f"vbvn_prev_{os.getpid()}{audio.suffix}"
            r = subprocess.run(
                [ffmpeg, "-i", str(audio), "-t", str(preview_seconds),
                 "-c", "copy", str(preview_audio), "-y", "-loglevel", "error"],
                capture_output=True, text=True,
            )
            if r.returncode == 0:
                audio = preview_audio

        dur_r = subprocess.run(
            [ffprobe, "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", str(audio)],
            capture_output=True, text=True,
        )
        duration = dur_r.stdout.strip()
        if not duration:
            log(f"Could not read duration of {audio.name}")
            return False
        duration_f = float(duration)

        # ── Transcription ───────────────────────────────────────────────────
        progress("Transcribing", 0)
        log(f"Transcribing {audio.name}…")

        ass_path = Path(tempfile.gettempdir()) / f"vbvn_cap_{os.getpid()}.ass"
        segments_path = Path(tempfile.gettempdir()) / f"vbvn_seg_{os.getpid()}.json"

        script = TRANSCRIBE_TEMPLATE.substitute(
            AUDIO=str(audio),
            CHUNK_SIZE=s.get("words_per_chunk", 2),
            WIDTH=width, HEIGHT=height,
            FONT_NAME=s.get("font_name", "Georgia"),
            FONT_SIZE=s.get("font_size", 64),
            CAPTION_HIGHLIGHT=s.get("caption_highlight", "&H004CA8C9"),
            CAPTION_COLOR=s.get("caption_color", "&H50D8EAF0"),
            CAPTION_BACK=s.get("caption_back", "&HBF000000"),
            CAPTION_Y=caption_y_ass,
            CAPTION_STYLE=s.get("caption_style", "karaoke"),
            ASS_PATH=str(ass_path),
            SEGMENTS_PATH=str(segments_path),
        )
        env = dict(os.environ)
        env["KMP_DUPLICATE_LIB_OK"] = "TRUE"
        rc, stderr = _run_transcription_with_progress(
            python_bin, env, script, duration_f,
            lambda frac: progress("Transcribing", frac))

        sub_filter = ""
        if rc != 0 or not ass_path.exists():
            log(f"Transcription failed:\n{stderr[-1000:]}")
        else:
            sub_filter = f",subtitles={ass_path.as_posix()}"

        # Save segments now so both outline and question cards can use them
        _saved_segments = []
        if segments_path.exists():
            try:
                _saved_segments = json.loads(segments_path.read_text())
            except Exception:
                pass
        segments_path.unlink(missing_ok=True)

        # ── Align questions FIRST so outline can be gapped at those times ───
        question_ass_path = None
        _q_drawbox = ""
        timed_q = []
        _q_blackout = []  # (start, end) intervals where question cards are shown

        if s.get("question_cards_enabled") and script_path and _saved_segments:
            progress("Building question cards", 0)
            q_raw = extract_discussion_questions(script_path)
            if q_raw:
                timed_q = align_questions_to_transcript(q_raw, _saved_segments)
                if timed_q:
                    for qi, (qs, _) in enumerate(timed_q):
                        qnext = timed_q[qi + 1][0] if qi + 1 < len(timed_q) else duration_f
                        _q_blackout.append((qs, min(qnext, qs + 55, duration_f)))
                    log(f"Discussion questions: {len(timed_q)} cards aligned.")
                else:
                    log("Could not align discussion questions to transcript.")
            else:
                log("No discussion questions found in script.")

        # ── Outline (gapped at question card timestamps) ─────────────────────
        if not outline_enabled:
            # No outline: finalise sidebar_w=0 and art_x now
            sidebar_w = 0
            art_x = _compute_art_x(0)

        if outline_enabled and _saved_segments:
            progress("Generating outline", 0)
            points = []

            if script_path:
                script_points = extract_outline_from_script(script_path)
                if script_points and outline_titles:
                    if isinstance(outline_titles[0], dict):
                        orig_map = {e["original"]: e["title"] for e in outline_titles}
                        script_points = [(orig_map[t], a) for t, a in script_points if t in orig_map]
                    else:
                        script_points = [
                            (outline_titles[i] if i < len(outline_titles) and outline_titles[i] else t, a)
                            for i, (t, a) in enumerate(script_points)
                        ]
                if script_points:
                    points = align_outline_to_transcript(script_points, _saved_segments)
                    log(f"Outline from script: {len(points)} sections aligned.")

            if not points:
                log("Generating outline via Ollama…")
                points = generate_outline(_saved_segments, s, log_cb=log)

            # Finalise sidebar_w and art_x based on whether points were found.
            # If no points, sidebar_w=0 so art is centred on the full screen.
            sidebar_w = _sidebar_w_req if points else 0
            art_x = _compute_art_x(sidebar_w)

            if points:
                outline_ass_path = Path(tempfile.gettempdir()) / f"vbvn_out_{os.getpid()}.ass"
                if outline_style == "ticker":
                    build_ticker_ass(points, s, duration_f, outline_ass_path)
                else:
                    outline_start_y = art_y + 16
                    outline_end_y   = height - wave_h - 12
                    build_sidebar_ass(points, s, duration_f, outline_ass_path,
                                      start_y=outline_start_y, end_y=outline_end_y,
                                      art_x=art_x, blackout=_q_blackout)
                sub_filter += f",subtitles={outline_ass_path.as_posix()}"
                log(f"Outline: {len(points)} points ({outline_style}).")
            else:
                log("No outline points — art centred, no sidebar reserved.")

        # ── Build question card ASS + drawbox ────────────────────────────────
        if timed_q:
            q_card_x, q_card_y = 0, 0
            q_card_w, q_card_h = width, height
            # Centre text on full screen (not just right panel — outline is hidden)
            q_text_cx = width // 2
            q_text_cy = (height - wave_h) // 2
            question_ass_path = Path(tempfile.gettempdir()) / f"vbvn_q_{os.getpid()}.ass"
            build_question_cards_ass(timed_q, s, duration_f, question_ass_path,
                                     q_card_x, q_card_y, q_card_w, q_card_h,
                                     text_cx=q_text_cx, text_cy=q_text_cy)
            sub_filter += f",subtitles={question_ass_path.as_posix()}"
            _q_drawbox = build_question_drawbox(timed_q, duration_f,
                                                q_card_x, q_card_y, q_card_w, q_card_h)

        # ── Title overlay ───────────────────────────────────────────────────
        title_filter = ""
        if title and s.get("title_enabled", True):
            escaped = title.replace("\\", "\\\\").replace("'", "\\'").replace(":", "\\:")
            fn = s.get("title_font", "Georgia").replace("'", "\\'")
            title_filter = (
                f",drawtext=text={escaped}:font='{fn}':"
                f"fontsize={s.get('title_font_size', 48)}:fontcolor={s.get('title_color', 'white')}:"
                f"x=(w-text_w)/2:y=72:box=1:boxcolor=black@0.4:boxborderw=10"
            )

        # ── Build ffmpeg filter_complex ──────────────────────────────────────
        progress("Rendering", 0)
        log(f"Rendering video…")

        tmp_audio = Path(tempfile.gettempdir()) / f"vbvn_audio_{os.getpid()}{audio.suffix}"
        shutil.copy(audio, tmp_audio)
        tmp_out = output.with_suffix(".tmp.mp4")

        bg_filter = build_bg_filter(s, fps)

        # Input layout: [0:v]=art (if enabled), [1:a]=audio (or [0:a] if no art)
        inputs = []
        audio_stream = "0:a"

        if art_enabled:
            inputs += ["-loop", "1", "-framerate", str(fps), "-i", str(art_path)]
            audio_stream = "1:a"

        inputs += ["-i", str(tmp_audio)]

        # QR code image input (optional)
        qr_enabled = s.get("qr_enabled") and s.get("qr_path") and Path(s["qr_path"]).exists()
        qr_idx = None
        if qr_enabled:
            qr_idx = 2 if art_enabled else 1
            inputs += ["-loop", "1", "-framerate", str(fps), "-i", str(s["qr_path"])]

        # Build filter graph
        # art input index is 0 if art_enabled, otherwise n/a
        # audio input index is 1 if art_enabled, else 0

        glow_sigma = max(1, s.get("glow_sigma", 40))
        fp = []  # filter_complex parts

        if art_enabled:
            audio_idx = 1
            fp.append(f"[0:v]scale={art_size}:{art_size},format=rgb24[art_sharp]")
            fp.append(bg_filter)
            if s.get("glow_enabled", True):
                fp.append(f"[art_sharp]gblur=sigma={glow_sigma}[glow_soft]")
                fp.append(f"[bg][glow_soft]overlay={art_x}:{art_y}:format=auto[bg_glow]")
                fp.append(f"[bg_glow][art_sharp]overlay={art_x}:{art_y}:format=auto[bg_art]")
            else:
                fp.append(f"[bg][art_sharp]overlay={art_x}:{art_y}:format=auto[bg_art]")
            prev = "[bg_art]"
        else:
            audio_idx = 0
            fp.append(bg_filter)
            prev = "[bg]"

        # Question card dark background covers the full screen, inserted BEFORE
        # the waveform so the waveform composites on top of it naturally.
        # Captions and QR are in later steps so they also show through.
        if _q_drawbox:
            fp.append(f"{prev}{_q_drawbox}[bg_qcard]")
            prev = "[bg_qcard]"

        if wave_enabled:
            wc_high = s.get("waveform_high_color", "#cc3333")  # peaks  → red
            wc_mid  = s.get("waveform_color",      "#c9a84c")  # normal → gold
            wc_low  = s.get("waveform_low_color",  "#3366cc")  # quiet  → blue
            yo = height - wave_h  # y offset for all wave overlays
            # Layer 1 (bottom): full volume → high color visible at peaks
            fp.append(f"[{audio_idx}:a]asplit=3[aw1][aw2][aw3]")
            fp.append(f"[aw1]showwaves=s={width}x{wave_h}:mode=cline:colors={wc_high},format=rgba[wv1]")
            # Layer 2: 55% volume → mid color covers the lower-amplitude portion
            fp.append(f"[aw2]volume=0.55,showwaves=s={width}x{wave_h}:mode=cline:colors={wc_mid},format=rgba[wv2]")
            # Layer 3: 20% volume → low color covers the quietest portion
            fp.append(f"[aw3]volume=0.20,showwaves=s={width}x{wave_h}:mode=cline:colors={wc_low},format=rgba[wv3]")
            fp.append(f"{prev}[wv1]overlay=0:{yo}:format=auto[bg_wv1]")
            fp.append(f"[bg_wv1][wv2]overlay=0:{yo}:format=auto[bg_wv2]")
            fp.append(f"[bg_wv2][wv3]overlay=0:{yo}:format=auto[bg_wave]")
            prev = "[bg_wave]"

        # Subtitles and title drawtext are chained as trailing filters on the last label
        extra = sub_filter.lstrip(",") + title_filter
        if extra:
            fp.append(f"{prev}{extra}[pre_qr]")
            prev = "[pre_qr]"
        # else prev stays as-is

        # QR code overlay in chosen corner, above the waveform strip
        if qr_enabled and qr_idx is not None:
            qr_size = s.get("qr_size", 160)
            qr_margin = 20
            corner = s.get("qr_corner", "bottom-right")
            qr_bottom = height - wave_h - qr_margin
            qr_y = qr_bottom - qr_size
            if "right" in corner:
                qr_x = width - qr_size - qr_margin
            else:
                qr_x = qr_margin
            if "top" in corner:
                qr_y = qr_margin
            fp.append(f"[{qr_idx}:v]scale={qr_size}:{qr_size}[qr_scaled]")
            fp.append(f"{prev}[qr_scaled]overlay={qr_x}:{qr_y}:format=auto[out]")
        elif extra:
            fp.append(f"[pre_qr]null[out]")
        else:
            fp.append(f"{prev}null[out]")

        filter_str = ";".join(fp)

        cmd = [
            ffmpeg,
            *inputs,
            "-filter_complex", filter_str,
            "-map", "[out]", "-map", f"{audio_idx}:a",
            "-c:v", "libx264", "-preset", "fast", "-crf", "18", "-pix_fmt", "yuv420p",
            "-af", "loudnorm=I=-16:TP=-1.5:LRA=11",
            "-c:a", "aac", "-b:a", "192k", "-ar", "44100",
            "-t", duration,
            "-movflags", "+faststart",
            "-loglevel", "warning",
            str(tmp_out), "-y",
        ]

        rc, stderr = _run_ffmpeg_with_progress(
            cmd, duration_f, lambda frac: progress("Rendering", frac))
        tmp_audio.unlink(missing_ok=True)

        if rc != 0:
            log(f"Render failed:\n{stderr[-2000:]}")
            tmp_out.unlink(missing_ok=True)
            return False

        shutil.move(str(tmp_out), str(output))
        size_mb = output.stat().st_size / (1024 * 1024)
        progress("Done", 1.0)
        log(f"Done → {output.name} ({size_mb:.1f} MB)")
        return True

    finally:
        if ass_path is not None:
            ass_path.unlink(missing_ok=True)
        if outline_ass_path is not None:
            outline_ass_path.unlink(missing_ok=True)
        if question_ass_path is not None:
            question_ass_path.unlink(missing_ok=True)
        if preview_audio is not None:
            preview_audio.unlink(missing_ok=True)
